from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SOURCE = Path(r"C:\Users\huday\OneDrive\Desktop\EXAMIFO_DESKTOP\Examifo_Implementation_and_Test_Status_Report.docx")
OUTPUT = Path(r"D:\maui\Examifo_Desktop\Examifo_Implementation_and_Test_Status_Report_with_Time_Log.docx")

doc = Document(SOURCE)
doc.add_page_break()
doc.add_heading("7. Work-session time log", level=1)

p = doc.add_paragraph()
p.add_run("Timing basis. ").bold = True
p.add_run(
    "The durations below are reconstructed from timestamps visible in the retained chat and screenshots. "
    "They represent elapsed collaboration windows, including explanation, implementation, builds, automated "
    "tests, manual testing, troubleshooting, and pauses for user feedback. They are not timesheet-grade coding "
    "hours. Where no reliable start/end pair exists, the entry is explicitly marked as unverified."
)

rows = [
    ("Portion 01 - Get Permission", "15 Aug 2026, 7:29 PM to about 11:15 PM", "Approx. 3 h 46 min",
     "High", "Point planning, implementation, automated suites, login/offline/restart/manual acceptance, and troubleshooting."),
    ("Portion 04 - Preserve Data", "15 Aug, about 11:20 PM to 16 Aug, 2:26 AM", "Approx. 3 h 06 min",
     "High", "Schema, encryption, persistence/outbox work, migrations, SQLite inspection, automated tests, and manual verification."),
    ("Responsive desktop window/UI adjustments", "16 Aug 2026, 3:08 AM to 3:15 AM", "Approx. 7 min",
     "High", "Window proportions, responsive XAML behavior, and Windows screen centering."),
    ("Portion 02 Points 1-5 - catalogue and temporary acquisition", "16 Aug 2026, visible evidence 1:16 PM to 1:25 PM", "At least 9 min; full span unavailable",
     "Medium", "Catalogue cache, metadata/manifest validation, safe download path, build/test work, and manual offline catalogue check."),
    ("Portion 02 Points 6-11 - verified package installation", "16 Aug 2026; no complete timestamp pair retained", "Not reliably measurable from chat",
     "Low", "SHA-256 verification, schema/protected-content checks, encrypted installation, atomic activation, rollback preservation, SQLite integration, tests, and build."),
]

table = doc.add_table(rows=1, cols=6)
table.style = "Table Grid"
table.autofit = False
headers = ["Work block", "Observed span", "Elapsed time", "Confidence", "Included work", "Classification"]
widths = [Inches(1.35), Inches(1.45), Inches(1.15), Inches(0.72), Inches(1.55), Inches(0.83)]
for i, (cell, header, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
    cell.width = width
    cell.text = header
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(255, 255, 255)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "1F4E78")
    cell._tc.get_or_add_tcPr().append(shading)

for task, span, elapsed, confidence, included in rows:
    cells = table.add_row().cells
    values = [task, span, elapsed, confidence, included, "Estimated"]
    for cell, value, width in zip(cells, values, widths):
        cell.width = width
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(8)

table.rows[-1].cells[-1].text = "Unverified"

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
r = p.add_run("Reconstructed total with measurable start/end pairs: approximately 6 hours 59 minutes, ")
r.bold = True
p.add_run(
    "plus the unmeasured portions of Portion 02. This total should be described as an evidence-based estimate, "
    "not a billing or productivity total."
)

doc.add_heading("Timing interpretation by milestone", level=2)
for text in [
    "Portion 01's fourteen milestones were handled within one continuous implementation-and-acceptance session; the retained chat does not support defensible point-by-point minute allocation.",
    "Portion 04's fourteen milestones were implemented and verified across the late-evening/early-morning session; SQLite inspection and manual troubleshooting account for a meaningful share of the elapsed window.",
    "Portion 02 Points 1-5 have only a narrow visible timestamp window, while Points 6-11 have test/build evidence but no reliable paired chat timestamps. Their exact elapsed durations remain unverified.",
]:
    doc.add_paragraph(text, style="List Bullet")

for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

doc.save(OUTPUT)
print(OUTPUT)
